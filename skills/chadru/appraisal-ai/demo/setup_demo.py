#!/usr/bin/env python3
"""
Demo Project Setup Script

Generates a complete set of demo files for testing the appraisal_ai pipeline:
  - Project folder with subject docs, comp docs, and exhibit docs (text files)
  - narrative.docx template (built with python-docx)
  - grid.xlsx template (built with openpyxl)
  - reference-data.yaml field mapping

Usage:
    ~/appraisal_ai/venv/bin/python demo/setup_demo.py

After running, test the pipeline with:
    /run demo/Demo-Project demo-land
"""

import os
import sys
import datetime

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DEMO_DIR = os.path.join(REPO_ROOT, 'demo')
PROJECT_DIR = os.path.join(DEMO_DIR, 'Demo-Project')
TEMPLATE_DIR = os.path.join(REPO_ROOT, 'templates', 'demo-land')

# ---------------------------------------------------------------------------
# Fictional data (all values are fake — no real properties or people)
# ---------------------------------------------------------------------------

SUBJECT = {
    'address': '1200 Maple Street',
    'city': 'Centerville',
    'state': 'Colorado',
    'zip': '80100',
    'owner': 'Greenfield Holdings, LLC',
    'parcel': '12-345-67-001',
    'land_sf': 25000,
    'land_acres': 0.57,
    'zoning': 'C-2: General Commercial',
    'flood': 'Zone X (unshaded)',
    'building_sf': 3200,
    'building_type': 'single-story retail building',
    'building_use': 'retail store',
    'building_year': 1985,
    'project_name': 'Maple Street Road Widening',
    'project_number': 'RDW M100-010',
    'project_code': '30100',
    'pe_id': 'PE-10',
    'pe_sf': 1500,
    'has_te': False,
    'te_id': 'TE-10',
    'te_sf': 400,
    'file_number': '2026-DEMO',
    'effective_date': 'January 15, 2026',
    'report_date': 'January 31, 2026',
    'psf_concluded': 50,
    'land_value': 1250000,
    'pe_value': 56250,
    'pe_percentage': 75,
    'client_contact': 'Jane Appraiser',
    'client_salutation': 'Ms. Appraiser',
    'client_email': 'jane.appraiser@example.com',
}

COMPS = [
    {
        'number': 1,
        'address': '800 Oak Avenue',
        'city': 'Centerville',
        'county': 'Demo',
        'state': 'CO',
        'zip': '80100',
        'price': 600000,
        'date': 'March 15, 2025',
        'date_short': '3/15/2025',
        'sf': 12000,
        'acres': 0.28,
        'psf': 50.00,
        'zoning': 'C-2',
        'flood': 'Zone X',
        'grantor': 'Maple Creek Partners',
        'grantee': 'Oak Street Ventures',
        'parcel': '12-345-67-010',
        'recording': '2025030100',
        'shape': 'Rectangular',
        'topography': 'Level',
        'utilities': 'All public',
        'comments': 'Vacant commercial land sold for future retail development. '
                    'Arms-length transaction confirmed with buyer.',
    },
    {
        'number': 2,
        'address': '2450 Pine Road',
        'city': 'Westfield',
        'county': 'Demo',
        'state': 'CO',
        'zip': '80105',
        'price': 1200000,
        'date': 'July 22, 2024',
        'date_short': '7/22/2024',
        'sf': 28000,
        'acres': 0.64,
        'psf': 42.86,
        'zoning': 'C-3',
        'flood': '40% in 100-year floodplain',
        'grantor': 'Sunrise Development Corp',
        'grantee': 'Westfield Town Center LLC',
        'parcel': '23-456-78-002',
        'recording': '2024071555',
        'shape': 'Irregular',
        'topography': 'Level to gently sloping',
        'utilities': 'All public',
        'comments': 'Large commercial parcel with partial flood zone exposure. '
                    'Buyer planned mixed-use development. Arms-length transaction.',
    },
    {
        'number': 3,
        'address': '550 Cedar Boulevard',
        'city': 'Centerville',
        'county': 'Demo',
        'state': 'CO',
        'zip': '80100',
        'price': 500000,
        'date': 'November 8, 2024',
        'date_short': '11/8/2024',
        'sf': 10500,
        'acres': 0.24,
        'psf': 47.62,
        'zoning': 'C-1',
        'flood': 'Zone X',
        'grantor': 'Cedar Valley Trust',
        'grantee': 'Centerville Retail Group',
        'parcel': '12-345-67-025',
        'recording': '2024110200',
        'shape': 'Rectangular',
        'topography': 'Level',
        'utilities': 'All public',
        'comments': 'Small commercial lot in established retail corridor. '
                    'Sold for land value. Arms-length transaction.',
    },
    {
        'number': 4,
        'address': '1800 Birch Lane',
        'city': 'Eastview',
        'county': 'Demo',
        'state': 'CO',
        'zip': '80110',
        'price': 750000,
        'date': 'May 3, 2025',
        'date_short': '5/3/2025',
        'sf': 15000,
        'acres': 0.34,
        'psf': 50.00,
        'zoning': 'C-2',
        'flood': 'Zone X',
        'grantor': 'Northern Star Realty',
        'grantee': 'Birch Holdings Inc',
        'parcel': '34-567-89-003',
        'recording': '2025050400',
        'shape': 'Rectangular',
        'topography': 'Level',
        'utilities': 'All public',
        'comments': 'Vacant commercial land at signalized intersection. '
                    'Purchased for retail pad development. Arms-length transaction.',
    },
]


# ===================================================================
# 1. Project folder text files
# ===================================================================

def write_subject_docs():
    """Write engagement letter, assessor card, and deed for the subject."""
    subject_dir = os.path.join(PROJECT_DIR, 'Subject')

    # Engagement letter
    with open(os.path.join(subject_dir, 'engagement_letter.txt'), 'w') as f:
        f.write(f"""ENGAGEMENT LETTER

Date: December 15, 2025

To: {SUBJECT['client_contact']}
    Colorado Department of Transportation
    Region 1 Right of Way
    2000 S Holly Street
    Denver, CO 80222

Re: Appraisal Services
    Project: {SUBJECT['project_name']} ({SUBJECT['project_number']})
    Property: {SUBJECT['address']}, {SUBJECT['city']}, {SUBJECT['state']} {SUBJECT['zip']}
    Owner: {SUBJECT['owner']}
    Parcel: {SUBJECT['parcel']}
    File Number: {SUBJECT['file_number']}

Dear {SUBJECT['client_salutation']}:

This letter confirms our engagement to prepare an appraisal of the above-referenced
property in connection with the {SUBJECT['project_name']} project. The purpose of the
appraisal is to estimate the market value of the property and the just compensation
for the proposed acquisition of a permanent easement ({SUBJECT['pe_id']}).

The permanent easement area contains approximately {SUBJECT['pe_sf']:,} square feet.

The effective date of value shall be January 15, 2026.

Sincerely,

Demo Appraisal Services, Inc.
""")

    # Assessor card
    with open(os.path.join(subject_dir, 'assessor_card.txt'), 'w') as f:
        f.write(f"""DEMO COUNTY ASSESSOR - PROPERTY RECORD CARD

Parcel Number:    {SUBJECT['parcel']}
Owner:            {SUBJECT['owner']}
Situs Address:    {SUBJECT['address']}
                  {SUBJECT['city']}, {SUBJECT['state']} {SUBJECT['zip']}

LAND DATA
  Land Area:      {SUBJECT['land_sf']:,} SF ({SUBJECT['land_acres']:.2f} acres)
  Zoning:         {SUBJECT['zoning']}
  Flood Zone:     {SUBJECT['flood']}
  Land Use Code:  Commercial

IMPROVEMENT DATA
  Year Built:     {SUBJECT['building_year']}
  Building Type:  {SUBJECT['building_type'].title()}
  Building SF:    {SUBJECT['building_sf']:,}
  Stories:        1
  Construction:   Wood frame with stucco exterior
  Roof:           Flat built-up
  Condition:      Average

ASSESSMENT (2025)
  Land Value:     $375,000
  Improvement:    $125,000
  Total Assessed: $500,000

LEGAL DESCRIPTION
  Lot 10, Block 3, Centerville Commercial Park, County of Demo,
  State of Colorado
""")

    # Deed
    with open(os.path.join(subject_dir, 'deed.txt'), 'w') as f:
        f.write(f"""RECORDING INFORMATION
Reception No.: 2018045200
Recorded: June 15, 2018
Demo County Clerk and Recorder

WARRANTY DEED

THIS DEED, made this 10th day of June, 2018, between

GRANTOR:  Centerville Properties LLC, a Colorado limited liability company

GRANTEE:  {SUBJECT['owner']}, a Colorado limited liability company

for the consideration of Ten Dollars and other good and valuable consideration,
the receipt and sufficiency of which are hereby acknowledged, hereby sells and
conveys to the Grantee the following described real property situated in
Demo County, State of Colorado:

Lot 10, Block 3, Centerville Commercial Park, according to the plat recorded
in Book 45 at Page 200, records of Demo County, Colorado.

Known and numbered as {SUBJECT['address']}, {SUBJECT['city']}, {SUBJECT['state']} {SUBJECT['zip']}

County Assessor Parcel Number: {SUBJECT['parcel']}

TOGETHER with all improvements, tenements, hereditaments, and appurtenances
thereto belonging, or in anywise appertaining.

IN WITNESS WHEREOF, the Grantor has executed this deed.

Centerville Properties LLC
By: Robert Chen, Manager
""")

    print(f"  Subject docs: 3 files written to {subject_dir}")


def write_comp_docs():
    """Write CoStar reports and deeds for each comparable sale."""
    comps_dir = os.path.join(PROJECT_DIR, 'Comparables')

    for comp in COMPS:
        folder = os.path.join(comps_dir, f"Sale {comp['number']} - {comp['address'].split(',')[0]}")
        os.makedirs(folder, exist_ok=True)

        # CoStar report
        with open(os.path.join(folder, 'costar_report.txt'), 'w') as f:
            f.write(f"""COSTAR COMPS - SALE DETAIL REPORT

Property:       {comp['address']}
                {comp['city']}, {comp['state']} {comp['zip']}
County:         {comp['county']}

SALE INFORMATION
  Sale Date:    {comp['date']}
  Sale Price:   ${comp['price']:,}
  Buyer:        {comp['grantee']}
  Seller:       {comp['grantor']}
  Document No.: {comp['recording']}

PROPERTY DETAILS
  Land Area:    {comp['sf']:,} SF ({comp['acres']:.2f} Acres)
  Price/SF:     ${comp['psf']:.2f}
  Zoning:       {comp['zoning']}
  Parcel:       {comp['parcel']}

SITE CHARACTERISTICS
  Shape:        {comp['shape']}
  Topography:   {comp['topography']}
  Utilities:    {comp['utilities']}
  Flood Zone:   {comp['flood']}

COMMENTS
  {comp['comments']}

VERIFICATION
  Confirmed arms-length transaction. No unusual conditions of sale noted.
  Price reflects land value only; no contributory improvement value.
""")

        # Deed
        with open(os.path.join(folder, 'deed.txt'), 'w') as f:
            f.write(f"""RECORDING INFORMATION
Reception No.: {comp['recording']}
Recorded: {comp['date']}
{comp['county']} County Clerk and Recorder

SPECIAL WARRANTY DEED

THIS DEED, made this day of {comp['date']}, between

GRANTOR:  {comp['grantor']}

GRANTEE:  {comp['grantee']}

for the consideration of ${comp['price']:,}, the receipt and sufficiency of
which are hereby acknowledged, hereby sells, conveys, and warrants to the
Grantee the following described real property situated in {comp['county']}
County, State of Colorado:

{comp['address']}, {comp['city']}, {comp['state']} {comp['zip']}

County Assessor Parcel Number: {comp['parcel']}

Land Area: {comp['sf']:,} SF ({comp['acres']:.2f} Acres)

IN WITNESS WHEREOF, the Grantor has executed this deed.

{comp['grantor']}
""")

    print(f"  Comp docs: {len(COMPS) * 2} files written to {comps_dir}")


def write_exhibit_docs():
    """Write flood map and zoning description exhibit files."""
    exhibits_dir = os.path.join(PROJECT_DIR, 'Exhibits')

    with open(os.path.join(exhibits_dir, 'flood_map.txt'), 'w') as f:
        f.write(f"""FEMA FLOOD MAP INFORMATION

Community:      City of Centerville, Demo County, Colorado
Map Panel:      08999C0150F
Effective Date: March 15, 2020

SUBJECT PROPERTY
  Address:      {SUBJECT['address']}, {SUBJECT['city']}, {SUBJECT['state']}
  Parcel:       {SUBJECT['parcel']}
  Flood Zone:   {SUBJECT['flood']}
  In SFHA:      No

Zone X (unshaded) — Areas determined to be outside the 0.2% annual chance
floodplain. No mandatory flood insurance purchase requirement.

The subject property is NOT located within a Special Flood Hazard Area (SFHA).
No portion of the parcel is within Zone A, AE, AH, AO, or VE.
""")

    with open(os.path.join(exhibits_dir, 'zoning_description.txt'), 'w') as f:
        f.write(f"""CITY OF CENTERVILLE ZONING INFORMATION

Property:       {SUBJECT['address']}
Parcel:         {SUBJECT['parcel']}
Current Zoning: {SUBJECT['zoning']}

C-2: GENERAL COMMERCIAL DISTRICT

PURPOSE
  The C-2 district is intended to provide areas for a wide range of commercial
  uses including retail, office, service, and limited mixed-use development.

PERMITTED USES
  - Retail stores and shops
  - Restaurants and food service
  - Professional and business offices
  - Personal service establishments
  - Financial institutions
  - Medical and dental offices
  - Hotels and motels (conditional use)

DEVELOPMENT STANDARDS
  Minimum Lot Size:     5,000 SF
  Minimum Lot Width:    50 feet
  Maximum Building Height: 45 feet (3 stories)
  Maximum Lot Coverage: 75%
  Front Setback:        10 feet
  Side Setback:         0 feet (5 feet adjacent to residential)
  Rear Setback:         10 feet
  Parking:              Per use requirements

NOTES
  The subject property's {SUBJECT['land_sf']:,} SF ({SUBJECT['land_acres']:.2f} acres) exceeds the
  minimum lot size of 5,000 SF. The existing {SUBJECT['building_type']} is a
  conforming use under C-2 zoning.
""")

    print(f"  Exhibit docs: 2 files written to {exhibits_dir}")


# ===================================================================
# 2. Template narrative.docx (built with python-docx)
# ===================================================================

def create_narrative_docx():
    """Build a demo appraisal narrative template using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_heading_centered(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return h

    def add_para(text, bold=False, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if alignment:
            p.alignment = alignment
        return p

    # ---------------------------------------------------------------
    # Cover Page
    # ---------------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    add_heading_centered('APPRAISAL REPORT', level=1)
    doc.add_paragraph()
    add_para('1200 Maple Street', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('Centerville, CO 80100', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('File Number: 2026-DEMO', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('Effective Date of Value: January 15, 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('Prepared for:', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('Colorado Department of Transportation', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('Prepared by:', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('Demo Appraisal Services, Inc.', alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Letter of Transmittal
    # ---------------------------------------------------------------
    add_heading_centered('LETTER OF TRANSMITTAL', level=1)
    doc.add_paragraph()
    add_para('January 31, 2026')
    doc.add_paragraph()
    add_para('Jane Appraiser')
    add_para('Colorado Department of Transportation')
    add_para('Region 1 Right of Way')
    add_para('2000 S Holly Street')
    add_para('Denver, CO 80222')
    doc.add_paragraph()
    add_para('Re: Appraisal of 1200 Maple Street, Centerville, CO 80100')
    add_para('    File Number: 2026-DEMO')
    add_para('    Project: Maple Street Road Widening (RDW M100-010)')
    add_para('    Owner: Greenfield Holdings, LLC')
    add_para('    Email: jane.appraiser@example.com')
    doc.add_paragraph()
    add_para('Dear Ms. Appraiser:')
    doc.add_paragraph()
    add_para(
        'In accordance with your request, I have prepared the accompanying '
        'appraisal report for the above-referenced property. The purpose of '
        'the appraisal is to estimate the market value of the property before '
        'the proposed acquisition and the just compensation for the permanent '
        'easement (PE-10) to be acquired by the Colorado Department of '
        'Transportation in connection with the Maple Street Road Widening '
        'project (RDW M100-010).'
    )
    doc.add_paragraph()
    add_para(
        'Based on my analysis and the data contained in this report, it is '
        'my opinion that the market value of the subject property, as of '
        'January 15, 2026, and subject to the assumptions and limiting '
        'conditions set forth in this report, is:'
    )
    doc.add_paragraph()
    add_para('Market Value Before the Taking: $1,250,000',
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('Just Compensation: $56,250',
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('Respectfully submitted,')
    add_para('Demo Appraisal Services, Inc.')

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Table of Contents
    # ---------------------------------------------------------------
    add_heading_centered('TABLE OF CONTENTS', level=1)
    doc.add_paragraph()
    toc_items = [
        'Letter of Transmittal',
        'Summary of Salient Facts',
        'Property Identification',
        'Purpose and Intended Use',
        'Scope of Work',
        'Property Description',
        'Highest and Best Use Analysis',
        'Sales Comparison Approach — Land Valuation',
        'LAND SALE MAP',
        'LAND SALE NO. 1',
        'LAND SALE NO. 2',
        'LAND SALE NO. 3',
        'LAND SALE NO. 4',
        'Comparative Analysis and Land Value Conclusion',
        'Permanent Easement Valuation',
        'Value of the Temporary Easement',
        'Allocation and Summary',
        'Assumptions and Limiting Conditions',
        'Appraiser Qualifications',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Summary of Salient Facts
    # ---------------------------------------------------------------
    add_heading_centered('SUMMARY OF SALIENT FACTS', level=1)
    doc.add_paragraph()

    table = doc.add_table(rows=18, cols=2)
    table.style = 'Table Grid'
    facts = [
        ('File Number', '2026-DEMO'),
        ('Property Address', '1200 Maple Street, Centerville, Colorado 80100'),
        ('Owner', 'Greenfield Holdings, LLC'),
        ('Parcel Number', '12-345-67-001'),
        ('Project', 'Maple Street Road Widening (RDW M100-010)'),
        ('Project Code', '30100'),
        ('Effective Date of Value', 'January 15, 2026'),
        ('Date of Report', 'January 31, 2026'),
        ('Land Area', '25,000 SF (0.57 Acres)'),
        ('Building', '3,200 SF single-story retail building (1985)'),
        ('Zoning', 'C-2: General Commercial'),
        ('Flood Zone', 'Zone X (unshaded)'),
        ('Highest and Best Use — Vacant', 'Commercial Development'),
        ('Highest and Best Use — Improved', 'Retail Commercial Building'),
        ('Permanent Easement (PE-10)', '1,500 SF'),
        ('Market Value Before Taking', '$1,250,000'),
        ('Value of PE-10', '$56,250'),
        ('Just Compensation', '$56,250'),
    ]
    for i, (label, value) in enumerate(facts):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Property Identification
    # ---------------------------------------------------------------
    add_heading_centered('PROPERTY IDENTIFICATION', level=2)
    doc.add_paragraph()
    add_para(
        'The subject property is located at 1200 Maple Street in the City '
        'of Centerville, Demo County, Colorado 80100. The property is '
        'identified by Demo County Assessor Parcel Number 12-345-67-001. '
        'The property is currently owned by Greenfield Holdings, LLC.'
    )
    doc.add_paragraph()
    add_para(
        'The subject is a 25,000 square foot (0.57 acre) commercially zoned '
        'parcel improved with a 3,200 square foot single-story retail building '
        'constructed in 1985. The building is of wood frame construction with '
        'stucco exterior walls, a flat built-up roof and adequate parking.'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Purpose and Intended Use
    # ---------------------------------------------------------------
    add_heading_centered('PURPOSE AND INTENDED USE', level=2)
    doc.add_paragraph()
    add_para(
        'The purpose of this appraisal is to estimate the market value of '
        'the subject property before the proposed acquisition and the just '
        'compensation for the permanent easement (PE-10) in connection with '
        'the Maple Street Road Widening project (RDW M100-010). The client '
        'is the Colorado Department of Transportation. The intended use is '
        'for right-of-way acquisition negotiations.'
    )

    # ---------------------------------------------------------------
    # Property Description (with traffic, frontage, surrounding)
    # ---------------------------------------------------------------
    add_heading_centered('PROPERTY DESCRIPTION', level=2)
    doc.add_paragraph()
    add_para('Site Description', bold=True)
    doc.add_paragraph()
    add_para(
        'The subject property is a 25,000 square foot (0.57 acre) parcel '
        'located on the south side of Maple Street between Oak Avenue and '
        'Cedar Boulevard in Centerville, Colorado.'
    )
    doc.add_paragraph()
    add_para(
        'Traffic Count: 18,500 vehicles/day Maple St (2025)'
    )
    add_para(
        'Frontage: Maple Street, Oak Avenue'
    )
    add_para(
        'Access: Maple Street and Oak Avenue'
    )
    doc.add_paragraph()
    add_para('Surrounding Development', bold=True)
    add_para('North: Maple Street / Commercial retail')
    add_para('South: Single-family residential')
    add_para('East: Oak Avenue / Mixed commercial')
    add_para('West: Cedar Boulevard / Office and retail')

    doc.add_paragraph()
    add_para('Utilities', bold=True)
    add_para('All public utilities are available to the site including water, sewer, natural gas, and electricity.')

    doc.add_paragraph()
    add_para('Zoning', bold=True)
    add_para(
        'The subject property is zoned C-2: General Commercial by the City '
        'of Centerville. The C-2 district permits a wide range of commercial '
        'uses including retail, office, and service establishments.'
    )

    doc.add_paragraph()
    add_para('Flood Zone', bold=True)
    add_para(
        'According to the FEMA Flood Insurance Rate Map, Community Panel '
        '08999C0150F dated March 15, 2020, the subject property is located '
        'in Zone X (unshaded), which is outside the Special Flood Hazard Area.'
    )

    doc.add_paragraph()
    add_para('Title Report', bold=True)
    add_para(
        'The title report dated September 10, 2025, by First Demo Title '
        'Insurance Company was reviewed. The report did not disclose any '
        'encumbrances that would adversely affect marketability or value.'
    )

    doc.add_paragraph()
    add_para('Improvement Description', bold=True)
    add_para(
        'The subject is improved with a 3,200 square foot single-story '
        'retail building constructed in 1985, currently used as a retail store. '
        'The building is of wood '
        'frame construction with stucco exterior walls, a flat built-up '
        'roof and adequate parking. The improvements are in average condition '
        'and contribute minimal value to the site.'
    )

    doc.add_paragraph()
    add_para('Property History', bold=True)
    add_para(
        'The subject property is owned by Greenfield Holdings, LLC. The '
        'property was acquired by the current owner on June 15, 2018 for '
        'a reported price of $400,000 from Centerville Properties LLC as '
        'recorded under Reception No. 2018045200.'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Easement Description
    # ---------------------------------------------------------------
    add_heading_centered('EASEMENT DESCRIPTION', level=2)
    doc.add_paragraph()
    add_para(
        'The Colorado Department of Transportation proposes to acquire a '
        'permanent easement (PE-10) containing approximately 1,500 square '
        'feet along the north side of the subject property in connection '
        'with the Maple Street Road Widening project (RDW M100-010).'
    )
    doc.add_paragraph()
    add_para(
        'The easement area is a strip along the Maple Street frontage '
        'and will be used for sidewalk improvements, utility relocation, '
        'and right-of-way widening. Greenfield Holdings, LLC is the grantor '
        'of the permanent easement.'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Highest and Best Use (duplicate value test)
    # ---------------------------------------------------------------
    add_heading_centered('HIGHEST AND BEST USE ANALYSIS', level=2)
    doc.add_paragraph()
    add_para(
        'Highest and best use is defined as the reasonably probable use of '
        'property that results in the highest value. The four criteria are '
        'legally permissible, physically possible, financially feasible, '
        'and maximally productive.'
    )
    doc.add_paragraph()

    # HBU table — both values are "Retail Commercial Building" in template
    # This tests the duplicate reference value handling
    hbu_table = doc.add_table(rows=3, cols=2)
    hbu_table.style = 'Table Grid'
    hbu_table.cell(0, 0).text = 'Highest and Best Use'
    hbu_table.cell(0, 1).text = 'Conclusion'
    hbu_table.cell(1, 0).text = 'As If Vacant'
    hbu_table.cell(1, 1).text = 'Retail Commercial Building'
    hbu_table.cell(2, 0).text = 'As Improved'
    hbu_table.cell(2, 1).text = 'Retail Commercial Building'

    doc.add_paragraph()
    add_para('As If Vacant', bold=True)
    add_para(
        'Considering the subject site as if vacant and available for '
        'development, the highest and best use is for commercial development '
        'consistent with the C-2 zoning. The site is physically capable of '
        'supporting commercial development, and such use would be financially '
        'feasible given current market conditions.'
    )
    doc.add_paragraph()
    add_para('As Improved', bold=True)
    add_para(
        'The property is currently improved with a Retail Commercial Building. '
        'The improvements contribute to the overall value of the property. '
        'The highest and best use as improved is the continued use as a '
        'Retail Commercial Building.'
    )
    doc.add_paragraph()
    add_para(
        'Improvements: Retail Commercial Building'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Sales Comparison Approach — Land Valuation
    # ---------------------------------------------------------------
    add_heading_centered('SALES COMPARISON APPROACH — LAND VALUATION', level=2)
    doc.add_paragraph()
    add_para(
        'In the Sales Comparison Approach, market value is estimated by '
        'comparing the subject to similar properties that have recently '
        'sold. Adjustments are made for differences in location, physical '
        'characteristics, and conditions of sale.'
    )
    doc.add_paragraph()
    add_para(
        'Four comparable land sales were identified in the Centerville and '
        'surrounding communities. The sales range from approximately 0.24 '
        'acres to 0.64 acres in size and sold between July 2024 and May '
        '2025. The unadjusted sale prices indicate a range from $42.86 PSF '
        'to $50.00 PSF.'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # LAND SALE MAP
    # ---------------------------------------------------------------
    add_heading_centered('LAND SALE MAP', level=2)
    doc.add_paragraph()
    add_para('[Land Sale Location Map — to be inserted]',
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(20):
        doc.add_paragraph()

    doc.add_page_break()

    # ---------------------------------------------------------------
    # LAND SALE NO. 1–4 (comp writeup sections)
    # ---------------------------------------------------------------
    for comp in COMPS:
        add_heading_centered(f"LAND SALE NO. {comp['number']}", level=2)
        doc.add_paragraph()

        detail_lines = [
            f"Location:\t{comp['address']}",
            f"City/County:\t{comp['city']}/{comp['county']}",
            f"Assessor Account No.:\t{comp['parcel']}",
            f"Grantor:\t{comp['grantor']}",
            f"Grantee:\t{comp['grantee']}",
            f"Recording/Reception No.:\t{comp['recording']}",
            f"Sale Date:\t{comp['date']}",
            f"Sale Price:\t${comp['price']:,}",
            f"Parcel Size:\t{comp['acres']:.2f} Acres ({comp['sf']:,} SF)",
            f"Unit Price:\t${comp['psf']:.2f} Per SF",
            f"Zoning:\t{comp['zoning']}",
            f"Comments:\t{comp['comments']}",
        ]
        for line in detail_lines:
            add_para(line)
            doc.add_paragraph()  # blank line between details

        doc.add_page_break()

        # Aerial placeholder
        add_heading_centered(f"LAND SALE NO. {comp['number']} AERIAL", level=2)
        doc.add_paragraph()
        add_para(f'[Aerial photograph of Sale {comp["number"]} — to be inserted]',
                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for _ in range(20):
            doc.add_paragraph()

        doc.add_page_break()

    # ---------------------------------------------------------------
    # Comparative Analysis and Land Value Conclusion
    # ---------------------------------------------------------------
    add_heading_centered('Comparative Analysis and Land Value Conclusion', level=2)
    doc.add_paragraph()
    add_para(
        'The four comparable sales were analyzed and adjusted for differences '
        'relative to the subject property. Adjustments were considered for '
        'market conditions (time), location, zoning, flood zone, and size.'
    )
    doc.add_paragraph()
    add_para(
        'The comparable sales range from approximately 0.24 acres to 0.64 '
        'acres in size. The sales occurred between July 2024 and May 2025, '
        'representing a time range from approximately 8.43 months prior to '
        'the effective date of value to 17.77 months prior to the date of value.'
    )
    doc.add_paragraph()

    # Comparable Sales After Adjustments table
    adj_table = doc.add_table(rows=6, cols=2)
    adj_table.style = 'Table Grid'
    adj_table.cell(0, 0).text = 'Comparable Sales After Adjustments'
    adj_table.cell(0, 1).text = ''
    adj_table.cell(1, 0).text = 'Sale No.'
    adj_table.cell(1, 1).text = 'Adjusted $/SF'
    adj_table.cell(2, 0).text = '1'
    adj_table.cell(2, 1).text = ' $       51 '
    adj_table.cell(3, 0).text = '2'
    adj_table.cell(3, 1).text = ' $       45 '
    adj_table.cell(4, 0).text = '3'
    adj_table.cell(4, 1).text = ' $       48 '
    adj_table.cell(5, 0).text = '4'
    adj_table.cell(5, 1).text = ' $       51 '

    doc.add_paragraph()
    add_para(
        'After all adjustments, the comparable sales indicate adjusted '
        'unit prices that range from $45 PSF to $51 PSF. The arithmetic '
        'average is $49 PSF. The median is $50 PSF.'
    )
    doc.add_paragraph()
    add_para(
        'Based on the analysis, I have concluded a unit value of $50.00 per '
        'square foot for the subject property.'
    )
    doc.add_paragraph()

    # Land value calculation
    calc_table = doc.add_table(rows=3, cols=2)
    calc_table.style = 'Table Grid'
    calc_table.cell(0, 0).text = 'Land Area'
    calc_table.cell(0, 1).text = '25,000 SF'
    calc_table.cell(1, 0).text = 'Unit Value'
    calc_table.cell(1, 1).text = '$50.00 PSF'
    calc_table.cell(2, 0).text = 'Market Value Before Taking'
    calc_table.cell(2, 1).text = '$1,250,000'

    doc.add_paragraph()
    add_para('25,000 SF x $50.00 PSF = $1,250,000', bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Permanent Easement Valuation
    # ---------------------------------------------------------------
    add_heading_centered('PERMANENT EASEMENT VALUATION', level=2)
    doc.add_paragraph()
    add_para(
        'The permanent easement (PE-10) contains approximately 1,500 square '
        'feet along the north side of the subject property fronting Maple '
        'Street. The easement will be used for sidewalk improvements, '
        'utility relocation, and right-of-way widening.'
    )
    doc.add_paragraph()
    add_para(
        'The degree of impact on the easement area equates to 75% of the '
        'full fee value. The easement area retains limited utility for the '
        'property owner but the primary use shifts to the public right-of-way.'
    )
    doc.add_paragraph()
    add_para(
        '1,500 SF x $50 PSF x 75% = $56,250',
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    # PE calculation table
    pe_table = doc.add_table(rows=4, cols=2)
    pe_table.style = 'Table Grid'
    pe_table.cell(0, 0).text = 'Permanent Easement (PE-10)'
    pe_table.cell(0, 1).text = ''
    pe_table.cell(1, 0).text = 'Area'
    pe_table.cell(1, 1).text = '1,500 SF'
    pe_table.cell(2, 0).text = 'Calculation'
    pe_table.cell(2, 1).text = '1,500 SF x $50 PSF x 75%'
    pe_table.cell(3, 0).text = 'PE Value'
    pe_table.cell(3, 1).text = '$56,250'

    doc.add_paragraph()
    add_para('Remainder Analysis', bold=True)
    add_para(
        'The remainder of the subject property after the permanent easement '
        'acquisition contains approximately 23,500 square feet. The '
        'remainder retains adequate size, shape, access, and utility for '
        'continued commercial use. No damages to the remainder are concluded.'
    )
    doc.add_paragraph()
    add_para(
        'Value of Remainder Before Taking: $1,193,750'
    )
    add_para(
        'Value of Remainder After Taking: $1,193,750'
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Value of the Temporary Easement (TE sections — for removal test)
    # ---------------------------------------------------------------
    add_heading_centered('VALUE OF THE TEMPORARY EASEMENT', level=2)
    doc.add_paragraph()
    add_para(
        'The temporary construction easement (TE-10) contains approximately '
        '400 square feet adjacent to the permanent easement area. The '
        'temporary easement is anticipated to be needed for a period of '
        'twelve (12) months during construction.'
    )
    doc.add_paragraph()
    add_para(
        'The term of the TCE is twelve months. Upon the TCE expiration, '
        'the area reverts to the property owner in its original condition.'
    )
    doc.add_paragraph()
    add_para('Temporary Easement Language', bold=True)
    add_para(
        'The grantor covenants and agrees that no building, structure, or '
        'obstruction will be placed, erected, or permitted to remain upon '
        'the TCE area during the construction period.'
    )
    doc.add_paragraph()
    add_para('No Interference', bold=True)
    add_para(
        'The property owner shall not interfere with or obstruct the '
        "Department's use of the temporary easement area. No vehicle, "
        'equipment, or material shall be stored within the TCE boundaries '
        'during the construction period.'
    )
    doc.add_paragraph()
    add_para('Restoration', bold=True)
    add_para(
        'Upon completion of construction, the Department shall restore the '
        'TCE area to a condition substantially similar to its pre-construction '
        'state. The Department shall restore the TCE area including grading, '
        'paving, and landscaping as applicable.'
    )
    doc.add_paragraph()
    add_para(
        'The value of the temporary easement is based on a lost rent '
        'calculation using the concluded unit value.'
    )
    doc.add_paragraph()
    add_para(
        '400 SF x $50.00 PSF x 10% x 1 year = $2,000',
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    # TE table
    te_table = doc.add_table(rows=4, cols=2)
    te_table.style = 'Table Grid'
    te_table.cell(0, 0).text = 'Temporary Easement (TE-10)'
    te_table.cell(0, 1).text = ''
    te_table.cell(1, 0).text = 'Area'
    te_table.cell(1, 1).text = '400 SF'
    te_table.cell(2, 0).text = 'Calculation'
    te_table.cell(2, 1).text = '400 SF x $50 PSF x 10% x 12 months'
    te_table.cell(3, 0).text = 'TE Value'
    te_table.cell(3, 1).text = '$2,000'

    doc.add_page_break()

    # ---------------------------------------------------------------
    # ALLOCATION and SUMMARY tables
    # ---------------------------------------------------------------
    add_heading_centered('ALLOCATION AND SUMMARY', level=2)
    doc.add_paragraph()

    alloc_table = doc.add_table(rows=7, cols=2)
    alloc_table.style = 'Table Grid'
    alloc_table.cell(0, 0).text = 'ALLOCATION'
    alloc_table.cell(0, 1).text = ''
    alloc_table.cell(1, 0).text = 'Market Value Before Taking'
    alloc_table.cell(1, 1).text = '$1,250,000'
    alloc_table.cell(2, 0).text = 'Value of Permanent Easement (PE-10)'
    alloc_table.cell(2, 1).text = '$56,250'
    alloc_table.cell(3, 0).text = 'Value of Temporary Easement (TE-10)'
    alloc_table.cell(3, 1).text = '$2,000'
    alloc_table.cell(4, 0).text = 'Damages to Remainder'
    alloc_table.cell(4, 1).text = '$0'
    alloc_table.cell(5, 0).text = 'Benefits to Remainder'
    alloc_table.cell(5, 1).text = '$0'
    alloc_table.cell(6, 0).text = 'Just Compensation'
    alloc_table.cell(6, 1).text = '$56,250'

    doc.add_paragraph()

    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.cell(0, 0).text = 'SUMMARY'
    summary_table.cell(0, 1).text = ''
    summary_table.cell(1, 0).text = 'Permanent Easement (PE-10)'
    summary_table.cell(1, 1).text = '$56,250'
    summary_table.cell(2, 0).text = 'Temporary Easement (TE-10)'
    summary_table.cell(2, 1).text = '$2,000'
    summary_table.cell(3, 0).text = 'Damages'
    summary_table.cell(3, 1).text = '$0'
    summary_table.cell(4, 0).text = 'Total Just Compensation'
    summary_table.cell(4, 1).text = '$56,250'

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Assumptions and Limiting Conditions
    # ---------------------------------------------------------------
    add_heading_centered('ASSUMPTIONS AND LIMITING CONDITIONS', level=2)
    doc.add_paragraph()
    conditions = [
        'This appraisal is subject to the following assumptions and limiting conditions:',
        '1. The appraiser has no present or prospective interest in the property.',
        '2. The appraiser has no personal interest or bias with respect to the parties involved.',
        '3. Information furnished by others is assumed to be reliable.',
        '4. No survey of the property has been made by the appraiser.',
        '5. The property is appraised free and clear of any liens or encumbrances unless otherwise stated.',
    ]
    for c in conditions:
        add_para(c)

    # ---------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------
    output_path = os.path.join(TEMPLATE_DIR, 'narrative.docx')
    doc.save(output_path)
    print(f"  Narrative template: {output_path}")


# ===================================================================
# 3. Template grid.xlsx (built with openpyxl)
# ===================================================================

def create_grid_xlsx():
    """Build a demo adjustment grid template using openpyxl."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers

    wb = Workbook()

    # Rename default sheet to match expected name
    ws = wb.active
    ws.title = 'Original Land Sale Grid'

    # Column widths
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 18
    for col in ['C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 18

    # Styles
    header_font = Font(name='Arial', size=10, bold=True)
    data_font = Font(name='Arial', size=10)
    currency_fmt = '#,##0'
    pct_fmt = '0.00%'
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
    header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    def style_row(row, font=data_font, fill=None):
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            cell = ws[f'{col}{row}']
            cell.font = font
            cell.border = thin_border
            if fill:
                cell.fill = fill

    # ---------------------------------------------------------------
    # Row 1: Headers
    # ---------------------------------------------------------------
    headers = ['', 'SUBJECT', 'SALE 1', 'SALE 2', 'SALE 3', 'SALE 4']
    for i, h in enumerate(headers):
        cell = ws.cell(row=1, column=i+1, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')

    # ---------------------------------------------------------------
    # Row labels (column A)
    # ---------------------------------------------------------------
    row_labels = {
        2: 'Sale Price',
        3: 'Real Property Rights',
        4: 'Adjustment',
        5: 'Adjusted Price',
        6: 'Financing Terms',
        7: 'Type',
        8: 'Adjustment',
        9: 'Adj. Gross %',
        10: 'Conditions of Sale',
        11: 'Type',
        12: 'Adjustment',
        13: 'Cap Rate Adj',
        14: 'Post-Sale Expenditures',
        15: 'Type',
        16: 'Adjustment',
        17: '$/SF at Sale',
        18: 'Market Conditions',
        19: 'Date Adjustment',
        20: 'Time Period',
        21: 'Sale Date',
        22: '$/SF Time Adj',
        23: '$/SF Final',
        24: 'Physical Characteristics',
        25: 'Land SF',
        26: 'Land Value',
        27: '',
        28: 'Location',
        29: 'Location Adj %',
        30: '',
        31: 'Zoning',
        32: 'Zoning Adj %',
        33: '',
        34: 'Flood Zone',
        35: 'Flood Adj %',
        36: '',
        37: 'Land SF',
        38: 'Size Adj %',
        39: '',
        40: 'Net Physical Adj',
        41: 'Adjusted $/SF',
    }

    for row_num, label in row_labels.items():
        ws.cell(row=row_num, column=1, value=label)
        style_row(row_num)

    # Style header rows
    for row_num in [1]:
        style_row(row_num, font=header_font, fill=header_fill)

    # ---------------------------------------------------------------
    # Subject column (B) — template values
    # ---------------------------------------------------------------
    ws['B2'] = 'N/A'
    ws['B3'] = 'Fee Simple'
    ws['B7'] = 'Assumes Cash Equivalent'
    ws['B11'] = 'Assumes Market'
    ws['B15'] = 'Assumes None'
    ws['B21'] = datetime.datetime(2026, 1, 15)
    ws['B21'].number_format = 'MM/DD/YYYY'
    ws['B25'] = 25000
    ws['B25'].number_format = currency_fmt
    ws['B28'] = '1200 Maple Street, Centerville'
    ws['B29'] = 'N/A'
    ws['B31'] = 'C-2: General Commercial'
    ws['B34'] = 'Zone X (unshaded)'
    ws['B35'] = '—'
    ws['B37'] = 25000
    ws['B37'].number_format = currency_fmt
    ws['B38'] = 'N/A'

    # ---------------------------------------------------------------
    # Comp columns (C–F) — template values for 4 comps
    # ---------------------------------------------------------------
    comp_cols = ['C', 'D', 'E', 'F']
    comp_data_for_grid = [
        {
            'price': 600000,
            'rights': 'Fee Simple',
            'financing': 'Cash Equivalent',
            'conditions': 'Market',
            'expenditures': 'None known',
            'date': datetime.datetime(2025, 3, 15),
            'sf': 12000,
            'location': '800 Oak Ave, Centerville',
            'zoning': 'C-2',
            'flood': 'Zone X',
        },
        {
            'price': 1200000,
            'rights': 'Fee Simple',
            'financing': 'Cash Equivalent',
            'conditions': 'Market',
            'expenditures': 'None known',
            'date': datetime.datetime(2024, 7, 22),
            'sf': 28000,
            'location': '2450 Pine Rd, Westfield',
            'zoning': 'C-3',
            'flood': '40% in 100-year floodplain',
        },
        {
            'price': 500000,
            'rights': 'Fee Simple',
            'financing': 'Cash Equivalent',
            'conditions': 'Market',
            'expenditures': 'None known',
            'date': datetime.datetime(2024, 11, 8),
            'sf': 10500,
            'location': '550 Cedar Blvd, Centerville',
            'zoning': 'C-1',
            'flood': 'Zone X',
        },
        {
            'price': 750000,
            'rights': 'Fee Simple',
            'financing': 'Cash Equivalent',
            'conditions': 'Market',
            'expenditures': 'None known',
            'date': datetime.datetime(2025, 5, 3),
            'sf': 15000,
            'location': '1800 Birch Ln, Eastview',
            'zoning': 'C-2',
            'flood': 'Zone X',
        },
    ]

    for idx, comp in enumerate(comp_data_for_grid):
        col = comp_cols[idx]

        # Sale Price (row 2)
        ws[f'{col}2'] = comp['price']
        ws[f'{col}2'].number_format = currency_fmt

        # Real Property Rights (row 3)
        ws[f'{col}3'] = comp['rights']

        # Adjustment row 4 = 0 (no rights adjustment)
        ws[f'{col}4'] = 0
        ws[f'{col}4'].number_format = pct_fmt

        # Row 5: Adjusted Price = Sale Price * (1 + row 4)
        ws[f'{col}5'] = f'={col}2*(1+{col}4)'
        ws[f'{col}5'].number_format = currency_fmt

        # Financing (row 7)
        ws[f'{col}7'] = comp['financing']

        # Financing adjustment (row 8)
        ws[f'{col}8'] = 0
        ws[f'{col}8'].number_format = pct_fmt

        # Row 9: Adj Gross % = row 4 + row 8
        ws[f'{col}9'] = f'={col}4+{col}8'
        ws[f'{col}9'].number_format = pct_fmt

        # Conditions (row 11)
        ws[f'{col}11'] = comp['conditions']

        # Conditions adjustment (row 12)
        ws[f'{col}12'] = 0
        ws[f'{col}12'].number_format = pct_fmt

        # Row 13: Cap Rate Adj = row 5 * (1 + row 12)
        ws[f'{col}13'] = f'={col}5*(1+{col}12)'
        ws[f'{col}13'].number_format = currency_fmt

        # Expenditures (row 15)
        ws[f'{col}15'] = comp['expenditures']

        # Expenditures adjustment (row 16)
        ws[f'{col}16'] = 0
        ws[f'{col}16'].number_format = currency_fmt

        # Row 17: $/SF at Sale = (row 13 + row 16) / row 25
        ws[f'{col}17'] = f'=({col}13+{col}16)/{col}25'
        ws[f'{col}17'].number_format = '#,##0.00'

        # Market conditions adjustment (row 19) — monthly rate
        ws[f'{col}19'] = 0.0006
        ws[f'{col}19'].number_format = '0.0000'

        # Sale Date (row 21)
        ws[f'{col}21'] = comp['date']
        ws[f'{col}21'].number_format = 'MM/DD/YYYY'

        # Row 22: $/SF Time Adj = row 17 * (1 + row 19 * months)
        # months = (B21 - col21) / 30.42
        ws[f'{col}22'] = f'={col}17*(1+{col}19*((B$21-{col}21)/30.42))'
        ws[f'{col}22'].number_format = '#,##0.00'

        # Row 23: $/SF Final (before physical adj) = row 22
        ws[f'{col}23'] = f'={col}22'
        ws[f'{col}23'].number_format = '#,##0.00'

        # Land SF (row 25)
        ws[f'{col}25'] = comp['sf']
        ws[f'{col}25'].number_format = currency_fmt

        # Row 26: Land Value = row 23 * row 25
        ws[f'{col}26'] = f'={col}23*{col}25'
        ws[f'{col}26'].number_format = currency_fmt

        # Location (row 28)
        ws[f'{col}28'] = comp['location']

        # Location Adj % (row 29) — APPRAISER JUDGMENT, leave at 0
        ws[f'{col}29'] = 0
        ws[f'{col}29'].number_format = pct_fmt

        # Zoning (row 31)
        ws[f'{col}31'] = comp['zoning']

        # Zoning Adj % (row 32) — APPRAISER JUDGMENT, leave at 0
        ws[f'{col}32'] = 0
        ws[f'{col}32'].number_format = pct_fmt

        # Flood Zone (row 34)
        ws[f'{col}34'] = comp['flood']

        # Flood Adj % (row 35) — APPRAISER JUDGMENT, leave at 0
        ws[f'{col}35'] = 0
        ws[f'{col}35'].number_format = pct_fmt

        # Land SF (row 37) — duplicate of row 25
        ws[f'{col}37'] = comp['sf']
        ws[f'{col}37'].number_format = currency_fmt

        # Size Adj % (row 38) — APPRAISER JUDGMENT, leave at 0
        ws[f'{col}38'] = 0
        ws[f'{col}38'].number_format = pct_fmt

        # Row 40: Net Physical Adj = row 29 + row 32 + row 35 + row 38
        ws[f'{col}40'] = f'={col}29+{col}32+{col}35+{col}38'
        ws[f'{col}40'].number_format = pct_fmt

        # Row 41: Adjusted $/SF = row 23 * (1 + row 40)
        ws[f'{col}41'] = f'={col}23*(1+{col}40)'
        ws[f'{col}41'].number_format = '#,##0.00'

    # ---------------------------------------------------------------
    # Add a second sheet (Sorted Summary) to test multi-sheet handling
    # ---------------------------------------------------------------
    ws2 = wb.create_sheet('Sorted Summary')
    ws2['A1'] = 'This sheet contains a sorted summary of the land sales.'
    ws2['A2'] = 'It references the Original Land Sale Grid sheet.'
    ws2['A3'] = 'Do not modify this sheet directly.'

    # ---------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------
    output_path = os.path.join(TEMPLATE_DIR, 'grid.xlsx')
    wb.save(output_path)
    print(f"  Grid template: {output_path}")


# ===================================================================
# 4. Reference data YAML
# ===================================================================

def create_reference_data_yaml():
    """Write the field mapping for the demo template."""
    sys.path.insert(0, os.path.join(REPO_ROOT, 'scripts'))
    from utils import save_yaml

    ref_data = {
        'template_type': 'demo-land',
        'has_te': True,  # Template has TE content (project may remove it)

        'fields': {
            # Owner
            'owner_name': 'Greenfield Holdings, LLC',

            # Client
            'client_contact': 'Jane Appraiser',
            'client_salutation': 'Ms. Appraiser',
            'client_email': 'jane.appraiser@example.com',
            'client_mailing_address': '1200 Maple Street',
            'client_mailing_csz': 'Centerville, Colorado 80100',

            # Address forms
            'address_forms': [
                '1200 Maple Street',
            ],

            'city_state_zip_full': 'Centerville, Colorado 80100',
            'city_state_zip_short': 'Centerville, CO 80100',

            # Parcel
            'parcel_id': '12-345-67-001',

            # File number
            'file_number': '2026-DEMO',

            # Dates
            'effective_date': 'January 15, 2026',
            'report_date': 'January 31, 2026',

            # Land
            'land_sf': '25,000',
            'land_acres': '0.57',

            # Building
            'building_sf': '3,200',
            'building_type': 'single-story retail building',
            'building_use': 'retail store',

            # Easements
            'pe_id': 'PE-10',
            'pe_sf': '1,500',
            'te_id': 'TE-10',
            'te_sf': '400',

            # Values
            'psf_concluded_2': '$50.00 PSF',
            'psf_concluded_3': '$50.00 per square foot',
            'land_value_total': '1,250,000',
            'before_take_value': '1,250,000',
            'pe_value': '56,250',
            'te_value': '2,000',
            'just_comp_total': '56,250',
            'remainder_before': '1,193,750',
            'remainder_after': '1,193,750',
            'pe_calculation': '1,500 SF x $50 PSF x 75%',
            'pe_percentage_calc': '75%',
            'te_rate': '10%',
            'te_term': '12 months',
            'te_term_spelled': 'twelve (12) months',

            # Zoning
            'zoning': 'C-2: General Commercial',

            # Project
            'project_name': 'Maple Street Road Widening',
            'project_number': 'RDW M100-010',
            'project_code': '30100',

            # Property history
            'property_history': 'The subject property is owned by Greenfield Holdings, LLC',

            # HBU — DUPLICATE VALUE TEST
            # Both set to "Retail Commercial Building" in template
            # Tests the duplicate-value special case handling
            'hbu_as_vacant': 'Retail Commercial Building',
            'hbu_as_improved': 'Retail Commercial Building',

            # Traffic / Frontage / Access
            'traffic_count': '18,500 vehicles/day Maple St (2025)',
            'frontage_line': 'Maple Street, Oak Avenue',
            'access_line': 'Maple Street and Oak Avenue',

            # Surrounding development
            'surrounding_east': 'Oak Avenue / Mixed commercial',
            'surrounding_west': 'Cedar Boulevard / Office and retail',

            # Title report
            'title_report': 'The title report dated September 10, 2025, by First Demo Title Insurance Company was reviewed.',

            # Building description
            'building_year_narrative': 'constructed in 1985',
            'building_construction_narrative': 'wood frame construction with stucco exterior walls',
            'building_roof_narrative': 'a flat built-up roof and adequate parking',

            # Easement grantor
            'easement_grantor': 'Greenfield Holdings, LLC',

            # PE percentage
            'pe_percentage_context': 'equates to 75% of the full fee value',

            # Land calculation
            'land_calc_psf': '$50.00 PSF',

            # Comp analysis
            'comp_range': 'from approximately 0.24 acres to 0.64 acres',
            'comp_date_range': 'between July 2024 and May 2025',
            'comp_time_range': 'from approximately 8.43 months prior to the effective date of value to 17.77 months prior to the date of value',

            # Comp adjusted PSF summary (table cells are rebuilt, not field-replaced;
            # these are kept for cross-reference verification only)
            'comp_adj_psf_low': ' $       45 ',
            'comp_adj_psf_high': ' $       51 ',
            'comp_psf_summary': 'adjusted unit prices that range from $45 PSF to $51 PSF. The arithmetic average is $49 PSF. The median is $50 PSF.',

            # Template-specific sections to remove
            'remove_sections': [
                '2018045200',
                'Centerville Properties LLC',
            ],
        },
    }

    output_path = os.path.join(TEMPLATE_DIR, 'reference-data.yaml')
    save_yaml(ref_data, output_path)
    print(f"  Reference data: {output_path}")


# ===================================================================
# Main
# ===================================================================

def main():
    print('Setting up demo project...\n')

    # Ensure directories exist
    for d in [
        os.path.join(PROJECT_DIR, 'Subject'),
        os.path.join(PROJECT_DIR, 'Comparables'),
        os.path.join(PROJECT_DIR, 'Exhibits'),
        os.path.join(PROJECT_DIR, 'Narrative'),
        TEMPLATE_DIR,
    ]:
        os.makedirs(d, exist_ok=True)

    # Create comp subdirectories
    for comp in COMPS:
        addr = comp['address'].split(',')[0]
        os.makedirs(
            os.path.join(PROJECT_DIR, 'Comparables', f"Sale {comp['number']} - {addr}"),
            exist_ok=True,
        )

    # 1. Project text files
    print('1. Writing project documents...')
    write_subject_docs()
    write_comp_docs()
    write_exhibit_docs()

    # 2. Template narrative
    print('\n2. Building narrative template...')
    create_narrative_docx()

    # 3. Template grid
    print('\n3. Building grid template...')
    create_grid_xlsx()

    # 4. Reference data
    print('\n4. Writing reference data...')
    create_reference_data_yaml()

    # Summary
    print('\n' + '=' * 60)
    print('Demo setup complete!')
    print('=' * 60)
    print(f'\nProject folder:  {PROJECT_DIR}')
    print(f'Template folder: {TEMPLATE_DIR}')
    print(f'\nTo test the pipeline:')
    print(f'  /run {PROJECT_DIR} demo-land')
    print(f'\nKey test scenarios:')
    print(f'  - has_te: false (tests TE section removal)')
    print(f'  - Comp 2 has flood zone (tests flood adjustment)')
    print(f'  - HBU duplicate values (tests special case handling)')
    print(f'  - 4 comps with deeds + CoStar (tests deed-first workflow)')


if __name__ == '__main__':
    main()
